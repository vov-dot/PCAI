"""
Модуль для работы с документами (Word, Excel)
"""
import os
from docx import Document
from docx.shared import Inches, Pt
import pandas as pd
from datetime import datetime


def create_word_document(filename: str, title: str, content: str) -> str:
    """
    Создает новый Word документ с заголовком и содержанием.
    
    Args:
        filename: Имя файла (без расширения .docx)
        title: Заголовок документа
        content: Основное содержание документа
    
    Returns:
        Путь к созданному файлу
    """
    if not filename.endswith('.docx'):
        filename = f"{filename}.docx"
    
    filepath = os.path.join(os.getcwd(), 'output', filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    doc = Document()
    
    # Добавляем заголовок
    heading = doc.add_heading(title, 0)
    
    # Добавляем содержание
    doc.add_paragraph(content)
    
    # Добавляем дату создания
    doc.add_paragraph(f"Создано: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    doc.save(filepath)
    
    return filepath


def create_excel_table(filename: str, data: list, columns: list = None) -> str:
    """
    Создает Excel таблицу из данных.
    
    Args:
        filename: Имя файла (без расширения .xlsx)
        data: Список строк данных (каждая строка - список значений)
        columns: Список названий колонок (опционально)
    
    Returns:
        Путь к созданному файлу
    """
    if not filename.endswith('.xlsx'):
        filename = f"{filename}.xlsx"
    
    filepath = os.path.join(os.getcwd(), 'output', filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    # Создаем DataFrame
    if columns:
        df = pd.DataFrame(data, columns=columns)
    else:
        df = pd.DataFrame(data)
    
    # Сохраняем в Excel
    df.to_excel(filepath, index=False, sheet_name='Data')
    
    return filepath


def append_to_word_document(filename: str, text: str) -> str:
    """
    Добавляет текст в существующий Word документ.
    
    Args:
        filename: Имя файла документа
        text: Текст для добавления
    
    Returns:
        Путь к обновленному файлу
    """
    if not filename.endswith('.docx'):
        filename = f"{filename}.docx"
    
    filepath = os.path.join(os.getcwd(), 'output', filename)
    
    if not os.path.exists(filepath):
        return create_word_document(filename, "Новый документ", text)
    
    doc = Document(filepath)
    doc.add_paragraph(text)
    doc.save(filepath)
    
    return filepath


def read_excel_file(filename: str) -> dict:
    """
    Читает Excel файл и возвращает данные.
    
    Args:
        filename: Имя файла для чтения
    
    Returns:
        Словарь с данными таблицы
    """
    if not filename.endswith('.xlsx'):
        filename = f"{filename}.xlsx"
    
    filepath = os.path.join(os.getcwd(), 'output', filename)
    
    if not os.path.exists(filepath):
        return {"error": f"Файл {filepath} не найден"}
    
    df = pd.read_excel(filepath)
    
    return {
        "columns": df.columns.tolist(),
        "data": df.values.tolist(),
        "shape": df.shape
    }
